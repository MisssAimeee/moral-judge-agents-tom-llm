#!/usr/bin/env python3
"""
export_prompts_docx.py  --  Standing exporter: write EVERY prompt template used in
the behavioral pipeline to a Word document, with the ORIGINAL human-study prompt
listed first. ANALYSIS/DOCS ONLY — no model inference.

The single source of truth is 03_behavioral.py (SCALE_QUESTIONS, SOURCE_SCALES,
PARA_TEMPLATES). Re-run this whenever the prompt set changes to keep the shareable
Word doc in sync:

    python code/export_prompts_docx.py

Output:
    outputs/prompts/ToM_prompts.docx
"""
import os, importlib.util, datetime

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.join(HERE, "..")


def load_behavioral():
    """Import 03_behavioral.py by path (its name starts with a digit)."""
    path = os.path.join(HERE, "03_behavioral.py")
    spec = importlib.util.spec_from_file_location("behavioral_mod", path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


# Which source papers use each human_verbatim scale (for provenance in the doc).
SCALE_SOURCES = {
    "permissibility": "Young, Cushman, Hauser & Saxe (2007, PNAS) — YS2008 stimuli "
                      "(the original belief×outcome moral-judgment study)",
    "blame":          "Young & Saxe (2009) — YS2009 stimuli",
    "wrongness":      "Young & Saxe (2011); Koster-Hale et al.; Saxe & Wexler; "
                      "Saxe & Kanwisher; Bruneau et al. — YS2011/KPH/SW/SK/BP stimuli",
}
# Order human_verbatim scales with the original (permissibility) first.
SCALE_ORDER = ["permissibility", "blame", "wrongness"]


def main():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    beh = load_behavioral()
    out_dir = os.path.join(ROOT, "outputs", "prompts")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "ToM_prompts.docx")

    doc = Document()

    title = doc.add_heading("ToM Project — Moral-Judgment Prompt Set", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"Auto-generated from code/03_behavioral.py on "
        f"{datetime.date.today().isoformat()}. "
        "Every model reads a vignette, then is asked one rating question. "
        "The story text is prepended; the question templates below are appended "
        "verbatim (separated by a blank line). All raw scales are normalized to a "
        "common 0–1 blameworthiness axis for cross-model/human comparison.")
    for run in sub.runs:
        run.italic = True

    # -------- Section 1: original human-study prompts (human_verbatim) --------
    doc.add_heading("1. Original human-study prompts (template: human_verbatim)", level=1)
    doc.add_paragraph(
        "These reproduce the exact rating scale of each source paper, auto-selected "
        "per stimulus by its source. Listed original-study-first.")
    for scale in SCALE_ORDER:
        q = beh.SCALE_QUESTIONS.get(scale)
        if not q:
            continue
        h = doc.add_heading(f"1.{SCALE_ORDER.index(scale)+1}  {scale} scale", level=2)
        src = doc.add_paragraph()
        src.add_run("Source: ").bold = True
        src.add_run(SCALE_SOURCES.get(scale, "—"))
        p = doc.add_paragraph(q.replace("{agent}", "<agent>"))
        p.paragraph_format.left_indent = Pt(18)
        for r in p.runs:
            r.font.name = "Consolas"; r.font.size = Pt(10)

    # map of which sources use which scale
    doc.add_heading("1.4  Source → scale mapping", level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "source key", "scale", "range"
    for key, (smin, smax, qkey) in beh.SOURCE_SCALES.items():
        c = tbl.add_row().cells
        c[0].text, c[1].text, c[2].text = key, qkey, f"{smin}–{smax}"

    # -------- Section 2: paraphrase / robustness templates --------
    doc.add_heading("2. Paraphrase & robustness templates", level=1)
    doc.add_paragraph(
        "Same vignettes, reworded questions, used to test that the intent-vs-outcome "
        "result is not a wording artifact (prompt-invariance analysis). The 3 most "
        "diagnostic (human_verbatim, para_wrong7, punish7) form the default run set.")
    for i, (name, (q, smin, smax)) in enumerate(beh.PARA_TEMPLATES.items(), 1):
        doc.add_heading(f"2.{i}  {name}   (scale {smin}–{smax})", level=2)
        p = doc.add_paragraph(q)
        p.paragraph_format.left_indent = Pt(18)
        for r in p.runs:
            r.font.name = "Consolas"; r.font.size = Pt(10)

    # -------- Section 3: full template list --------
    doc.add_heading("3. All template IDs", level=1)
    doc.add_paragraph("human_verbatim  (original, source-paper scales — see §1)")
    for name in beh.PARA_TEMPLATES:
        doc.add_paragraph(name, style="List Bullet")

    doc.save(out_path)
    print(f"wrote {os.path.relpath(out_path, ROOT)}")
    print(f"  templates exported: human_verbatim + {len(beh.PARA_TEMPLATES)} paraphrases "
          f"({', '.join(beh.PARA_TEMPLATES)})")


if __name__ == "__main__":
    main()
